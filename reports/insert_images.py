"""
向论文docx中插入补充图片
"""
import sys
sys.path.insert(0, r'F:\FILES\PROJECTS\PycharmProjects\dobot_magician\.venv\Lib\site-packages')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc_path = r'F:\FILES\PROJECTS\PycharmProjects\dobot_magician\reports\基于GMM-GMR运动基元的Dobot Magician夹爪码垛轨迹学习仿真研究.docx'
img_dir = r'F:\FILES\PROJECTS\PycharmProjects\dobot_magician\reports\images'

doc = Document(doc_path)

def add_image_with_caption(doc, img_path, caption_text, width_inches=5.5):
    """添加图片并附带图注"""
    # 添加图片
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    
    # 添加图注
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(caption_text)
    run.font.size = Pt(10.5)
    run.font.bold = False
    
    return caption

# 遍历段落，找到插入位置
paragraphs = list(doc.paragraphs)
insert_points = {}

for i, para in enumerate(paragraphs):
    text = para.text.strip()
    # 找到各章节标题
    if '2.1' in text and '软件架构' in text:
        insert_points['system_arch'] = i
    elif '2.3' in text and '示教数据' in text:
        insert_points['task_stages'] = i
    elif '2.4' in text and '训练' in text and '流程' in text:
        insert_points['flowchart'] = i
    elif '3.1' in text and 'GMM' in text and 'GMR' in text:
        insert_points['gmm'] = i
    elif '4.1' in text and '实验' in text:
        insert_points['3d_traj'] = i
    elif '5.1' in text or ('5.2' not in text and '5.' in text and '仿真' in text):
        if '5.1' in text:
            insert_points['sim_scene'] = i

print("找到插入位置:", insert_points)

# 我们采用在文档末尾添加所有图片的方式，更简单可靠
# 先添加分页符
doc.add_page_break()

# 添加补充图片章节标题
heading = doc.add_heading('补充图表', level=1)

# 1. 系统架构图
doc.add_heading('图1 系统软件架构图', level=2)
add_image_with_caption(doc, 
    os.path.join(img_dir, 'system_architecture.png'),
    '图1 项目软件架构与模块划分')

# 2. 算法流程图
doc.add_heading('图2 算法训练与回放流程图', level=2)
add_image_with_caption(doc,
    os.path.join(img_dir, 'algorithm_flow.png'),
    '图2 示教数据训练、评估与仿真回放完整流程')

# 3. 任务阶段图
doc.add_heading('图3 码垛任务阶段示意图', level=2)
add_image_with_caption(doc,
    os.path.join(img_dir, 'task_stages.png'),
    '图3 单点码垛任务的八个关键动作阶段')

# 4. GMM分量示意图
doc.add_heading('图4 GMM高斯混合分量示意图', level=2)
add_image_with_caption(doc,
    os.path.join(img_dir, 'gmm_components.png'),
    '图4 高斯混合模型(GMM)多分量分布可视化')

# 5. 3D轨迹图
doc.add_heading('图5 三维空间轨迹示意图', level=2)
add_image_with_caption(doc,
    os.path.join(img_dir, '3d_trajectory.png'),
    '图5 码垛任务在三维空间中的闭合运动轨迹')

# 6. 仿真场景图
doc.add_heading('图6 CoppeliaSim仿真场景图', level=2)
add_image_with_caption(doc,
    os.path.join(img_dir, 'simulation_scene.png'),
    '图6 CoppeliaSim简化夹爪码垛仿真场景')

# 保存
output_path = r'F:\FILES\PROJECTS\PycharmProjects\dobot_magician\reports\基于GMM-GMR运动基元的Dobot Magician夹爪码垛轨迹学习仿真研究_补充图片版.docx'
doc.save(output_path)
print(f"保存成功: {output_path}")
